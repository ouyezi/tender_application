"""Extract high-fidelity tables from DOCX/PDF and normalize image links.

Table extraction is best-effort per table: a single failing table is
recorded as a warning and skipped rather than aborting the whole file (see
docs/superpowers/specs/2026-07-16-workspace-management-design.md §5.2).
"""

from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from typing import Any

from docx import Document

IMAGE_LINK_RE = re.compile(r"(!\[[^\]]*\]\()([^)\s]+)(\))")


def normalize_image_paths(markdown: str, file_id: str) -> str:
    """Rewrite markdown image links to the artifact-relative ``../image/{file_id}/...`` form.

    Links already pointing at ``../image/...`` or at an http(s) URL are left
    untouched; anything else (bare filenames, absolute paths from
    conversion) is rewritten to just its basename under the file's image dir.
    """

    def repl(m: re.Match[str]) -> str:
        prefix, target, suffix = m.group(1), m.group(2), m.group(3)
        if target.startswith("../image/") or target.startswith(("http://", "https://")):
            return m.group(0)
        name = Path(target).name
        return f"{prefix}../image/{file_id}/{name}{suffix}"

    return IMAGE_LINK_RE.sub(repl, markdown)


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _table_to_html(table: Any) -> str:
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    if n_rows == 0 or n_cols == 0:
        return "<table></table>\n"

    grid = [[table.cell(r, c) for c in range(n_cols)] for r in range(n_rows)]
    visited: set[int] = set()
    row_htmls: list[str] = []

    for r in range(n_rows):
        cell_htmls: list[str] = []
        for c in range(n_cols):
            cell = grid[r][c]
            key = id(cell._tc)
            if key in visited:
                continue

            colspan = 1
            while c + colspan < n_cols and id(grid[r][c + colspan]._tc) == key:
                colspan += 1

            rowspan = 1
            while r + rowspan < n_rows and id(grid[r + rowspan][c]._tc) == key:
                rowspan += 1

            for rr in range(r, r + rowspan):
                for cc in range(c, c + colspan):
                    visited.add(id(grid[rr][cc]._tc))

            attrs = ""
            if rowspan > 1:
                attrs += f' rowspan="{rowspan}"'
            if colspan > 1:
                attrs += f' colspan="{colspan}"'
            cell_htmls.append(f"<td{attrs}>{_escape_html(cell.text)}</td>")

        row_htmls.append("<tr>" + "".join(cell_htmls) + "</tr>")

    return "<table>\n" + "\n".join(row_htmls) + "\n</table>\n"


def _table_to_csv_best_effort(table: Any) -> str | None:
    """Flatten a table to CSV; returns ``None`` if any row's cell count is irregular.

    Merged cells make a rectangular CSV lossy (spans collapse to a single
    repeated value per grid cell), so this is only attempted when the table
    has a uniform column count across all rows.
    """
    n_cols = len(table.columns)
    rows: list[list[str]] = []
    for row in table.rows:
        cells = row.cells
        if len(cells) != n_cols:
            return None
        rows.append([cell.text for cell in cells])

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerows(rows)
    return buf.getvalue()


def extract_tables_from_docx(path: str | Path, out_dir: str | Path) -> tuple[list[str], list[str]]:
    """Extract every top-level table in a DOCX to ``out_dir`` as HTML (+ CSV when possible).

    Returns ``(table_ids, warnings)``. Numbering (``tbl_001``, ``tbl_002``, ...)
    matches the placeholders emitted by ``convert.convert_docx_to_markdown``.
    A single table's failure is recorded as a warning and does not raise.
    """
    path = Path(path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    warnings: list[str] = []
    table_ids: list[str] = []

    try:
        document = Document(str(path))
        tables = document.tables
    except Exception as exc:  # pragma: no cover - defensive, whole-file open failure
        warnings.append(f"docx_open_failed: {exc}")
        return [], warnings

    for idx, table in enumerate(tables, start=1):
        tbl_id = f"tbl_{idx:03d}"
        try:
            html = _table_to_html(table)
            (out_dir / f"{tbl_id}.html").write_text(html, encoding="utf-8")
            csv_text = _table_to_csv_best_effort(table)
            if csv_text is not None:
                (out_dir / f"{tbl_id}.csv").write_text(csv_text, encoding="utf-8")
            table_ids.append(tbl_id)
        except Exception as exc:
            warnings.append(f"table_extract_failed:{tbl_id}:{exc}")

    return table_ids, warnings


def extract_tables_from_pdf(path: str | Path, out_dir: str | Path) -> tuple[list[str], list[str]]:
    """Best-effort table extraction from a PDF using PyMuPDF's layout analysis.

    Returns ``(table_ids, warnings)``. Merged-looking cells are approximated
    via PyMuPDF's ``find_tables`` (no true rowspan/colspan info is available
    from PDF layout analysis, so plain ``<td>`` cells are emitted).
    """
    import fitz  # PyMuPDF

    path = Path(path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    warnings: list[str] = []
    table_ids: list[str] = []

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        warnings.append(f"pdf_open_failed: {exc}")
        return [], warnings

    idx = 0
    try:
        for page in doc:
            try:
                found = page.find_tables()
            except Exception as exc:
                warnings.append(f"table_find_failed:page_{page.number}:{exc}")
                continue
            for table in found.tables:
                idx += 1
                tbl_id = f"tbl_{idx:03d}"
                try:
                    rows = table.extract()
                    html_rows = "\n".join(
                        "<tr>" + "".join(f"<td>{_escape_html(str(cell or ''))}</td>" for cell in row) + "</tr>"
                        for row in rows
                    )
                    (out_dir / f"{tbl_id}.html").write_text(
                        f"<table>\n{html_rows}\n</table>\n", encoding="utf-8"
                    )
                    buf = io.StringIO()
                    writer = csv.writer(buf)
                    writer.writerows([[cell or "" for cell in row] for row in rows])
                    (out_dir / f"{tbl_id}.csv").write_text(buf.getvalue(), encoding="utf-8")
                    table_ids.append(tbl_id)
                except Exception as exc:
                    warnings.append(f"table_extract_failed:{tbl_id}:{exc}")
    finally:
        doc.close()

    return table_ids, warnings
