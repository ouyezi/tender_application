from __future__ import annotations

import json
from collections import Counter
from pathlib import Path
from typing import Any, Optional

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker

from app import db as database
from app.config import REPORT_DIR
from app.models import DiagnosisResult
from app.services import artifact

COMPLIANCE_LABELS = {
    "satisfied": "满足",
    "violated": "违反",
    "cannot_satisfy": "无法满足",
    "insufficient_evidence": "证据不足",
    "manual_required": "需线下核验",
}

CONSEQUENCE_LABELS = {
    "no_score": "不得分",
    "bid_unusable": "投标无效",
    "score_risk": "得分风险",
    "general_risk": "一般风险",
}


def _result_label(item: dict[str, Any]) -> str:
    raw = item.get("compliance_status") or item.get("result", "")
    if raw in COMPLIANCE_LABELS:
        return COMPLIANCE_LABELS[raw]
    return str(raw)


def _format_consequence_tags(tags: Any) -> str:
    if not tags:
        return ""
    if isinstance(tags, str):
        try:
            tags = json.loads(tags)
        except json.JSONDecodeError:
            return tags
    if not isinstance(tags, list):
        return str(tags)
    labels = [
        CONSEQUENCE_LABELS.get(tag, tag)
        for tag in tags
        if isinstance(tag, str) and tag
    ]
    return "、".join(labels)


def build_markdown(task_id: str, results: list[dict]) -> str:
    total = len(results)
    counts = Counter(_result_label(item) for item in results)
    overview_parts = [f"总计 {total} 项"]
    for label in COMPLIANCE_LABELS.values():
        if counts.get(label):
            overview_parts.append(f"{label} {counts[label]} 项")
    for label, n in sorted(counts.items()):
        if label and label not in COMPLIANCE_LABELS.values():
            overview_parts.append(f"{label} {n} 项")

    lines = [
        "# 标书诊断报告",
        "",
        f"**任务编号：** {task_id}",
        "",
        "## 概览",
        "",
        "；".join(overview_parts),
        "",
        "## 诊断明细",
        "",
    ]

    for i, item in enumerate(results, start=1):
        title = item.get("content_title") or f"检查项 {i}"
        consequence_text = _format_consequence_tags(item.get("consequence_tags"))
        lines.extend(
            [
                f"### {i}. {title}",
                "",
                f"- **描述：** {item.get('description', '')}",
                f"- **结论：** {_result_label(item)}",
            ]
        )
        if consequence_text:
            lines.append(f"- **后果标签：** {consequence_text}")
        lines.extend(
            [
                f"- **证据：** {item.get('evidence', '')}",
                f"- **建议：** {item.get('suggestion', '')}",
                "",
            ]
        )

    return "\n".join(lines).rstrip() + "\n"


def write_docx(path: str, markdown: str) -> None:
    doc = Document()
    for line in markdown.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "":
            continue
        else:
            doc.add_paragraph(line)
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


async def generate_and_save_reports(
    task_id: str,
    session_factory: Optional[async_sessionmaker] = None,
) -> tuple[str, str]:
    factory = session_factory or database.SessionLocal
    async with factory() as session:
        result = await session.execute(
            select(DiagnosisResult)
            .where(DiagnosisResult.task_id == task_id)
            .order_by(DiagnosisResult.sort_order)
        )
        rows = list(result.scalars().all())

    results: list[dict[str, Any]] = []
    for row in rows:
        try:
            consequence_tags = json.loads(row.consequence_tags or "[]")
        except json.JSONDecodeError:
            consequence_tags = []
        results.append(
            {
                "content_title": row.content_title,
                "description": row.description,
                "result": row.result,
                "compliance_status": row.compliance_status,
                "consequence_tags": consequence_tags,
                "evidence": row.evidence,
                "suggestion": row.suggestion,
            }
        )

    out_dir = Path(REPORT_DIR) / task_id
    out_dir.mkdir(parents=True, exist_ok=True)
    md_path = out_dir / "report.md"
    docx_path = out_dir / "report.docx"

    markdown = build_markdown(task_id, results)
    md_path.write_text(markdown, encoding="utf-8")
    write_docx(str(docx_path), markdown)
    artifact.sync_to_artifact_report(task_id, md_path, docx_path)

    return str(md_path), str(docx_path)
