from __future__ import annotations

import io
import json

import pytest
from docx import Document
from sqlalchemy import select

from app import db as database
from app.models import ParseJob, WorkspaceFile
from app.services import artifact, parse_scheduler


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("总则", level=1)
    doc.add_paragraph("正文内容。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _noop_kick() -> None:
    return None


async def _create_task(client, monkeypatch) -> str:
    """Create a task via the tasks API with the (slow, real) parse background
    worker disabled, so files stay in ``parse_status="pending"`` and tests
    stay fast/deterministic. Individual tests write fake parse artifacts
    directly when they need a specific ``parse_status``/tree/markdown."""
    monkeypatch.setattr(parse_scheduler, "kick", _noop_kick)
    files = {
        "tender_file": (
            "tender.docx",
            io.BytesIO(_docx_bytes()),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "bid_file": (
            "bid.docx",
            io.BytesIO(_docx_bytes()),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    }
    r = await client.post(
        "/api/tasks", data={"background": "x", "requirements": "y"}, files=files
    )
    assert r.status_code == 201
    return r.json()["id"]


async def _write_fake_parse(task_id: str, file_id: str, markdown: str, tree: dict) -> None:
    root = artifact.artifact_root(task_id)
    md_dir = root / "markdown"
    json_dir = root / "json"
    md_dir.mkdir(parents=True, exist_ok=True)
    json_dir.mkdir(parents=True, exist_ok=True)
    md_path = md_dir / f"{file_id}.md"
    tree_path = json_dir / f"{file_id}.tree.json"
    md_path.write_text(markdown, encoding="utf-8")
    tree_path.write_text(json.dumps(tree, ensure_ascii=False), encoding="utf-8")

    async with database.SessionLocal() as session:
        wf = await session.get(WorkspaceFile, file_id)
        wf.md_path = str(md_path)
        wf.tree_path = str(tree_path)
        wf.parse_status = "succeeded"
        await session.commit()


def _sample_tree(markdown: str) -> dict:
    purpose_start = markdown.index("阐明")
    heading_start = markdown.index("## 目的")
    return {
        "nodes": [
            {
                "id": "n_root",
                "title": "总则",
                "level": 1,
                "numbering": None,
                "parent_id": None,
                "start_offset": 0,
                "end_offset": len(markdown),
                "self_start": 0,
                "subtree_end": len(markdown),
                "source": "heading",
                "children": [
                    {
                        "id": "n_child",
                        "title": "目的",
                        "level": 2,
                        "numbering": "1.1",
                        "parent_id": "n_root",
                        "start_offset": purpose_start,
                        "end_offset": len(markdown),
                        "self_start": heading_start,
                        "subtree_end": len(markdown),
                        "source": "heading",
                        "children": [],
                    }
                ],
            }
        ],
        "warnings": [],
    }


@pytest.mark.asyncio
async def test_list_and_detail_workspaces(client, monkeypatch):
    task_id = await _create_task(client, monkeypatch)

    r = await client.get("/api/workspaces")
    assert r.status_code == 200
    items = r.json()
    match = next(item for item in items if item["task_id"] == task_id)
    assert match["file_count"] == 2

    r2 = await client.get(f"/api/workspaces/{task_id}")
    assert r2.status_code == 200
    detail = r2.json()
    assert len(detail["files"]) == 2


@pytest.mark.asyncio
async def test_get_workspace_404_for_unknown_task(client):
    r = await client.get("/api/workspaces/does-not-exist")
    assert r.status_code == 404


@pytest.mark.asyncio
async def test_upload_document_file(client, monkeypatch):
    task_id = await _create_task(client, monkeypatch)

    r = await client.post(
        f"/api/workspaces/{task_id}/files",
        data={"label": "补充文件"},
        files={
            "file": (
                "extra.docx",
                io.BytesIO(_docx_bytes()),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 201
    body = r.json()
    assert body["kind"] == "document"
    assert body["parse_status"] in ("pending", "running")
    assert body["label"] == "补充文件"

    detail = (await client.get(f"/api/workspaces/{task_id}")).json()
    assert len(detail["files"]) == 3


@pytest.mark.asyncio
async def test_upload_other_extension_is_skipped(client, monkeypatch):
    task_id = await _create_task(client, monkeypatch)

    r = await client.post(
        f"/api/workspaces/{task_id}/files",
        data={"label": "备注"},
        files={"file": ("notes.txt", io.BytesIO(b"hello"), "text/plain")},
    )
    assert r.status_code == 201
    body = r.json()
    assert body["kind"] == "other"
    assert body["parse_status"] == "skipped"


@pytest.mark.asyncio
async def test_tree_and_content(client, monkeypatch):
    task_id = await _create_task(client, monkeypatch)
    detail = (await client.get(f"/api/workspaces/{task_id}")).json()
    file_id = detail["files"][0]["id"]

    markdown = "# 总则\n\n第一章内容。\n\n## 目的\n\n阐明目的。\n"
    tree = _sample_tree(markdown)
    await _write_fake_parse(task_id, file_id, markdown, tree)

    r = await client.get(f"/api/workspaces/{task_id}/files/{file_id}/tree")
    assert r.status_code == 200
    nodes = r.json()
    assert nodes[0]["id"] == "n_root"
    assert nodes[0]["numbering"] == ""  # normalized from None
    assert nodes[0]["children"][0]["id"] == "n_child"

    child = tree["nodes"][0]["children"][0]
    r2 = await client.get(
        f"/api/workspaces/{task_id}/files/{file_id}/content",
        params={"node_id": "n_child"},
    )
    assert r2.status_code == 200
    content = r2.json()
    assert content["title"] == "目的"
    assert content["markdown"] == markdown[child["start_offset"] : child["end_offset"]]


@pytest.mark.asyncio
async def test_invalid_node_id_returns_404(client, monkeypatch):
    task_id = await _create_task(client, monkeypatch)
    detail = (await client.get(f"/api/workspaces/{task_id}")).json()
    file_id = detail["files"][0]["id"]

    markdown = "# 总则\n\n第一章内容。\n\n## 目的\n\n阐明目的。\n"
    tree = _sample_tree(markdown)
    await _write_fake_parse(task_id, file_id, markdown, tree)

    r = await client.get(
        f"/api/workspaces/{task_id}/files/{file_id}/content",
        params={"node_id": "does-not-exist"},
    )
    assert r.status_code == 404


@pytest.mark.asyncio
async def test_reparse_failed_file(client, monkeypatch):
    task_id = await _create_task(client, monkeypatch)
    detail = (await client.get(f"/api/workspaces/{task_id}")).json()
    file_id = detail["files"][0]["id"]

    async with database.SessionLocal() as session:
        wf = await session.get(WorkspaceFile, file_id)
        wf.parse_status = "failed"
        wf.parse_error = "boom"
        await session.commit()

    r = await client.post(f"/api/workspaces/{task_id}/files/{file_id}/reparse")
    assert r.status_code == 200
    body = r.json()
    assert body["parse_status"] == "pending"
    assert body["parse_error"] is None

    async with database.SessionLocal() as session:
        jobs = (
            await session.execute(select(ParseJob).where(ParseJob.file_id == file_id))
        ).scalars().all()
        assert any(j.attempt >= 2 for j in jobs)

    r2 = await client.post(f"/api/workspaces/{task_id}/files/{file_id}/reparse")
    assert r2.status_code == 409


@pytest.mark.asyncio
async def test_download_and_index(client, monkeypatch):
    task_id = await _create_task(client, monkeypatch)
    detail = (await client.get(f"/api/workspaces/{task_id}")).json()
    file_id = detail["files"][0]["id"]

    r = await client.get(f"/api/workspaces/{task_id}/files/{file_id}/download")
    assert r.status_code == 200

    r2 = await client.get(f"/api/workspaces/{task_id}/index")
    assert r2.status_code == 200
    assert file_id in r2.text
