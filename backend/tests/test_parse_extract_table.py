from __future__ import annotations

from pathlib import Path

from docx import Document

from app.services.parse import extract


def _make_docx_with_merged_table(path: Path) -> None:
    doc = Document()
    doc.add_heading("表格示例", level=1)
    table = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"r{r}c{c}"
    # Horizontal merge (colspan) on row 0, columns 0-1.
    table.cell(0, 0).merge(table.cell(0, 1))
    # Vertical merge (rowspan) on column 2, rows 1-2.
    table.cell(1, 2).merge(table.cell(2, 2))
    doc.save(str(path))


def test_extract_tables_from_docx_preserves_merged_cells(tmp_path):
    docx_path = tmp_path / "sample.docx"
    _make_docx_with_merged_table(docx_path)
    out_dir = tmp_path / "table_out"

    table_ids, warnings = extract.extract_tables_from_docx(docx_path, out_dir)

    assert table_ids == ["tbl_001"]
    assert warnings == []

    html = (out_dir / "tbl_001.html").read_text(encoding="utf-8")
    assert "rowspan" in html or "colspan" in html
    assert "colspan" in html
    assert "rowspan" in html


def test_extract_tables_from_docx_single_table_failure_is_warned(tmp_path, monkeypatch):
    docx_path = tmp_path / "sample.docx"
    _make_docx_with_merged_table(docx_path)
    out_dir = tmp_path / "table_out"

    def _boom(_table):
        raise RuntimeError("boom")

    monkeypatch.setattr(extract, "_table_to_html", _boom)

    table_ids, warnings = extract.extract_tables_from_docx(docx_path, out_dir)

    assert table_ids == []
    assert len(warnings) == 1
    assert "tbl_001" in warnings[0]
