from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from app.services import artifact
from app.services.parse import pipeline


def _make_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("总则", level=1)
    doc.add_paragraph("本章描述项目背景与总体要求。")
    doc.add_heading("资格要求", level=2)
    doc.add_paragraph("投标人应具备相应资质。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "资质"
    table.cell(1, 1).text = "合格"
    doc.save(str(path))


async def test_run_parse_pipeline_docx_succeeds(tmp_path, monkeypatch):
    monkeypatch.setattr(artifact, "UPLOAD_DIR", tmp_path)

    task_id = "T-TEST-PIPE-001"
    file_id = "fdoc001"
    artifact.ensure_artifact_dirs(task_id)

    src_path = artifact.artifact_root(task_id) / "document" / f"{file_id}_sample.docx"
    src_path.parent.mkdir(parents=True, exist_ok=True)
    _make_sample_docx(src_path)

    result = await pipeline.run_parse_pipeline(file_id, task_id, str(src_path))

    assert result["status"] in {"succeeded", "partial"}
    assert result["error"] is None

    md_path = Path(result["md_path"])
    tree_path = Path(result["tree_path"])
    chunks_path = Path(result["chunks_path"])

    assert md_path.is_file()
    assert tree_path.is_file()
    assert chunks_path.is_file()

    markdown = md_path.read_text(encoding="utf-8")
    assert "总则" in markdown
    assert "table:tbl_001" in markdown

    tree = json.loads(tree_path.read_text(encoding="utf-8"))
    assert tree["nodes"]

    chunks = json.loads(chunks_path.read_text(encoding="utf-8"))
    assert len(chunks) >= 1
    assert all("chunk_id" in c and "node_id" in c and "title_path" in c for c in chunks)

    meta_path = artifact.artifact_root(task_id) / "json" / f"{file_id}.meta.json"
    assert meta_path.is_file()

    table_html = artifact.artifact_root(task_id) / "table" / file_id / "tbl_001.html"
    assert table_html.is_file()


async def test_run_parse_pipeline_unsupported_extension_fails(tmp_path, monkeypatch):
    monkeypatch.setattr(artifact, "UPLOAD_DIR", tmp_path)

    task_id = "T-TEST-PIPE-002"
    file_id = "fdoc002"
    artifact.ensure_artifact_dirs(task_id)

    src_path = artifact.artifact_root(task_id) / "document" / f"{file_id}_sample.txt"
    src_path.parent.mkdir(parents=True, exist_ok=True)
    src_path.write_text("plain text", encoding="utf-8")

    result = await pipeline.run_parse_pipeline(file_id, task_id, str(src_path))

    assert result["status"] == "failed"
    assert result["error"] is not None
