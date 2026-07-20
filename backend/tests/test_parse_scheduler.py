from __future__ import annotations

import asyncio
from pathlib import Path

import pytest
from docx import Document
from sqlalchemy import select

from app import db as database
from app.models import ParseJob, WorkspaceFile
from app.services import artifact, parse_scheduler, workspace
from app.services.parse.pipeline import run_parse_pipeline

TERMINAL_PARSE_STATUSES = {"succeeded", "partial", "failed"}


def _make_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("总则", level=1)
    doc.add_paragraph("本章描述项目背景与总体要求。")
    doc.save(str(path))


async def _wait_for_file_status(file_id: str, timeout: float = 10.0) -> WorkspaceFile:
    loop = asyncio.get_running_loop()
    deadline = loop.time() + timeout
    while True:
        async with database.SessionLocal() as session:
            wf = await session.get(WorkspaceFile, file_id)
            if wf is not None and wf.parse_status in TERMINAL_PARSE_STATUSES:
                return wf
        if loop.time() >= deadline:
            raise TimeoutError(
                f"file {file_id} did not reach a terminal parse_status within {timeout}s "
                f"(status={wf.parse_status if wf else 'missing'})"
            )
        await asyncio.sleep(0.05)


async def _create_file_and_job(task_id: str, file_id: str, stored_path: str) -> None:
    async with database.SessionLocal() as session:
        session.add(
            WorkspaceFile(
                id=file_id,
                task_id=task_id,
                label="招标文件",
                original_filename="sample.docx",
                stored_path=stored_path,
                kind="document",
                ext=".docx",
                parse_status="pending",
            )
        )
        session.add(
            ParseJob(
                file_id=file_id,
                task_id=task_id,
                status="queued",
                stage="convert",
            )
        )
        await session.commit()


@pytest.mark.asyncio
async def test_kick_runs_queued_job_end_to_end(tmp_path, monkeypatch, client):
    monkeypatch.setattr(artifact, "UPLOAD_DIR", tmp_path / "uploads")
    monkeypatch.setattr(workspace, "UPLOAD_DIR", tmp_path / "uploads")
    monkeypatch.setattr(
        "app.services.parse_scheduler.run_parse_pipeline",
        run_parse_pipeline,
    )

    task_id = "T-PARSE-SCHED-001"
    file_id = "fsched001"
    root = artifact.ensure_artifact_dirs(task_id)
    src_path = root / "document" / f"{file_id}_sample.docx"
    src_path.parent.mkdir(parents=True, exist_ok=True)
    _make_sample_docx(src_path)

    await _create_file_and_job(task_id, file_id, str(src_path))

    await parse_scheduler.kick()

    wf = await _wait_for_file_status(file_id)
    assert wf.parse_status in {"succeeded", "partial"}
    assert wf.md_path and Path(wf.md_path).is_file()
    assert wf.tree_path and Path(wf.tree_path).is_file()
    assert wf.chunks_path and Path(wf.chunks_path).is_file()

    async with database.SessionLocal() as session:
        job = (
            await session.execute(select(ParseJob).where(ParseJob.file_id == file_id))
        ).scalar_one()
        assert job.status == "succeeded"
        assert job.finished_at is not None
        assert job.started_at is not None

    index_path = root / "index.md"
    assert index_path.is_file()
    assert file_id in index_path.read_text(encoding="utf-8")


@pytest.mark.asyncio
async def test_kick_marks_job_failed_on_unsupported_extension(tmp_path, monkeypatch, client):
    monkeypatch.setattr(artifact, "UPLOAD_DIR", tmp_path / "uploads")
    monkeypatch.setattr(workspace, "UPLOAD_DIR", tmp_path / "uploads")
    monkeypatch.setattr(
        "app.services.parse_scheduler.run_parse_pipeline",
        run_parse_pipeline,
    )

    task_id = "T-PARSE-SCHED-002"
    file_id = "fsched002"
    root = artifact.ensure_artifact_dirs(task_id)
    src_path = root / "document" / f"{file_id}_sample.txt"
    src_path.parent.mkdir(parents=True, exist_ok=True)
    src_path.write_text("plain text", encoding="utf-8")

    await _create_file_and_job(task_id, file_id, str(src_path))

    await parse_scheduler.kick()

    wf = await _wait_for_file_status(file_id)
    assert wf.parse_status == "failed"
    assert wf.parse_error is not None

    async with database.SessionLocal() as session:
        job = (
            await session.execute(select(ParseJob).where(ParseJob.file_id == file_id))
        ).scalar_one()
        assert job.status == "failed"
        assert job.stage == "convert"


@pytest.mark.asyncio
async def test_kick_wiring_with_monkeypatched_pipeline(tmp_path, monkeypatch, client):
    """Unit test of scheduler wiring only: pipeline itself is mocked out."""
    task_id = "T-PARSE-SCHED-003"
    file_id = "fsched003"

    await _create_file_and_job(task_id, file_id, str(tmp_path / "unused.docx"))

    async def _fake_pipeline(file_id_arg, task_id_arg, stored_path_arg, **kwargs):
        assert file_id_arg == file_id
        assert task_id_arg == task_id
        return {
            "status": "succeeded",
            "md_path": str(tmp_path / "fake.md"),
            "tree_path": str(tmp_path / "fake.tree.json"),
            "chunks_path": str(tmp_path / "fake.chunks.json"),
            "error": None,
            "warnings": [],
        }

    monkeypatch.setattr(parse_scheduler, "run_parse_pipeline", _fake_pipeline)

    await parse_scheduler.kick()

    wf = await _wait_for_file_status(file_id)
    assert wf.parse_status == "succeeded"
    assert wf.md_path == str(tmp_path / "fake.md")


@pytest.mark.asyncio
async def test_recover_interrupted_parse_jobs_resets_running_to_queued(client):
    task_id = "T-PARSE-SCHED-004"
    file_id = "fsched004"

    async with database.SessionLocal() as session:
        session.add(
            WorkspaceFile(
                id=file_id,
                task_id=task_id,
                label="标书",
                original_filename="sample.docx",
                stored_path="/tmp/doesnotmatter.docx",
                kind="document",
                ext=".docx",
                parse_status="running",
            )
        )
        session.add(
            ParseJob(
                file_id=file_id,
                task_id=task_id,
                status="running",
                stage="chunk",
            )
        )
        await session.commit()

    await database.recover_interrupted_parse_jobs()

    async with database.SessionLocal() as session:
        wf = await session.get(WorkspaceFile, file_id)
        job = (
            await session.execute(select(ParseJob).where(ParseJob.file_id == file_id))
        ).scalar_one()
        assert wf.parse_status == "pending"
        assert job.status == "queued"
        assert job.stage == "convert"
